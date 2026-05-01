"""
Resume Parser Service
Extracts structured information from PDF and DOCX resumes using NLP.
"""

import re
import logging
from pathlib import Path
from typing import Optional
import io

logger = logging.getLogger(__name__)


class ResumeParserService:
    """
    Parses PDF/DOCX resumes into structured data.
    Uses pdfplumber for PDFs and python-docx for DOCX files.
    Applies regex + heuristic NLP for structured extraction.
    """

    # Common section headers
    SECTION_HEADERS = {
        "experience": ["experience", "work experience", "employment", "professional experience", "work history"],
        "education": ["education", "academic background", "qualifications", "academic qualifications"],
        "skills": ["skills", "technical skills", "core competencies", "technologies", "expertise"],
        "projects": ["projects", "personal projects", "key projects", "notable projects"],
        "certifications": ["certifications", "certificates", "licenses", "awards"],
        "summary": ["summary", "objective", "profile", "about me", "professional summary"],
    }

    # Known tech skills for extraction
    TECH_SKILLS = {
        # Languages
        "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust", "ruby", "php",
        "swift", "kotlin", "scala", "r", "matlab", "sql", "bash", "shell",
        # Frontend
        "react", "vue", "angular", "svelte", "next.js", "nuxt", "html", "css", "tailwind",
        "sass", "webpack", "vite", "redux", "graphql",
        # Backend
        "node.js", "fastapi", "django", "flask", "spring", "express", "laravel", "rails",
        # Databases
        "postgresql", "mysql", "mongodb", "redis", "sqlite", "elasticsearch", "dynamodb",
        "cassandra", "neo4j",
        # Cloud & DevOps
        "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ansible", "jenkins",
        "github actions", "ci/cd", "linux", "nginx",
        # AI/ML
        "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "keras", "hugging face",
        "langchain", "openai", "llm", "machine learning", "deep learning", "nlp", "computer vision",
        # Tools
        "git", "jira", "confluence", "figma", "postman", "rest api", "graphql",
        # Soft skills
        "leadership", "communication", "agile", "scrum", "problem solving",
    }

    def __init__(self):
        self._spacy_nlp = None  # lazy-loaded

    def _get_nlp(self):
        """Lazy-load spaCy model."""
        if self._spacy_nlp is None:
            try:
                import spacy
                self._spacy_nlp = spacy.load("en_core_web_sm")
            except Exception:
                logger.warning("spaCy model not available, using regex-only parsing")
                self._spacy_nlp = False
        return self._spacy_nlp if self._spacy_nlp is not False else None

    def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF bytes using pdfplumber."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            # Fallback to PyMuPDF
            try:
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                return "\n".join(page.get_text() for page in doc)
            except Exception as e2:
                logger.error(f"PyMuPDF fallback also failed: {e2}")
                raise ValueError(f"Could not extract text from PDF: {e2}")

    def extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX bytes using python-docx."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Could not extract text from DOCX: {e}")

    def extract_text(self, file_bytes: bytes, filename: str) -> str:
        """Route to correct extractor based on file extension."""
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return self.extract_text_from_pdf(file_bytes)
        elif ext in [".docx", ".doc"]:
            return self.extract_text_from_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Use PDF or DOCX.")

    # ─── Field Extraction ──────────────────────────────────────────────────

    def _extract_email(self, text: str) -> Optional[str]:
        pattern = r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
        match = re.search(pattern, text)
        return match.group() if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        pattern = r"(\+?\d[\d\s\-().]{7,14}\d)"
        match = re.search(pattern, text)
        return match.group().strip() if match else None

    def _extract_linkedin(self, text: str) -> Optional[str]:
        pattern = r"linkedin\.com/in/[\w\-]+"
        match = re.search(pattern, text, re.IGNORECASE)
        return f"https://{match.group()}" if match else None

    def _extract_github(self, text: str) -> Optional[str]:
        pattern = r"github\.com/[\w\-]+"
        match = re.search(pattern, text, re.IGNORECASE)
        return f"https://{match.group()}" if match else None

    def _extract_name(self, text: str) -> Optional[str]:
        """
        Heuristic: the name is usually in the first 3 lines,
        all-caps or title-case, no special chars.
        """
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines[:5]:
            # Skip lines that look like headers or contact info
            if re.search(r"[@|//|resume|cv]", line, re.IGNORECASE):
                continue
            # 2-4 words, title/upper case
            words = line.split()
            if 2 <= len(words) <= 4 and all(w.replace("-", "").replace(".", "").isalpha() for w in words):
                return line.title()
        return None

    def _extract_skills(self, text: str) -> list:
        """Match known tech skills against resume text."""
        text_lower = text.lower()
        found = set()
        for skill in self.TECH_SKILLS:
            # Use word boundary matching
            pattern = r"\b" + re.escape(skill) + r"\b"
            if re.search(pattern, text_lower):
                found.add(skill.title() if len(skill) > 3 else skill.upper())
        return sorted(list(found))

    def _split_sections(self, text: str) -> dict:
        """Split resume text into named sections."""
        lines = text.split("\n")
        sections = {"header": [], "other": []}
        current_section = "header"

        for line in lines:
            line_lower = line.strip().lower()
            matched_section = None

            for section_key, keywords in self.SECTION_HEADERS.items():
                if any(line_lower == kw or line_lower.startswith(kw) for kw in keywords):
                    if len(line.strip()) < 50:  # short lines are likely headers
                        matched_section = section_key
                        break

            if matched_section:
                current_section = matched_section
                if current_section not in sections:
                    sections[current_section] = []
            else:
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append(line)

        return {k: "\n".join(v) for k, v in sections.items()}

    def _parse_education(self, edu_text: str) -> list:
        """Parse education section into structured entries."""
        results = []
        degree_patterns = [
            r"(B\.?S\.?|B\.?E\.?|B\.?Tech|Bachelor[s]?|Master[s]?|M\.?S\.?|M\.?E\.?|M\.?Tech|Ph\.?D|MBA|B\.?A|M\.?A)",
        ]
        lines = [l.strip() for l in edu_text.split("\n") if l.strip()]
        i = 0
        while i < len(lines):
            for pattern in degree_patterns:
                if re.search(pattern, lines[i], re.IGNORECASE):
                    entry = {
                        "degree": lines[i],
                        "institution": lines[i + 1] if i + 1 < len(lines) else "",
                        "year": None,
                        "gpa": None,
                        "field": None,
                    }
                    # Try to find year
                    for j in range(i, min(i + 4, len(lines))):
                        year_match = re.search(r"\b(19|20)\d{2}\b", lines[j])
                        if year_match:
                            entry["year"] = year_match.group()
                        gpa_match = re.search(r"GPA[:\s]+(\d\.\d+)", lines[j], re.IGNORECASE)
                        if gpa_match:
                            entry["gpa"] = gpa_match.group(1)
                    results.append(entry)
                    break
            i += 1
        return results

    def _parse_experience(self, exp_text: str) -> list:
        """Parse work experience into structured entries."""
        results = []
        lines = [l.strip() for l in exp_text.split("\n") if l.strip()]

        date_pattern = re.compile(
            r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|"
            r"March|April|June|July|August|September|October|November|December)"
            r"[\s,]+\d{4}|Present|Current|\d{4}\s*[-–]\s*(\d{4}|Present|Current)",
            re.IGNORECASE,
        )

        current_entry = None
        for line in lines:
            if date_pattern.search(line):
                if current_entry:
                    results.append(current_entry)
                current_entry = {
                    "title": line,
                    "company": "",
                    "duration": date_pattern.search(line).group(),
                    "description": [],
                    "start_date": None,
                    "end_date": None,
                }
            elif current_entry:
                if not current_entry["company"] and not line.startswith("•") and not line.startswith("-"):
                    current_entry["company"] = line
                elif line.startswith(("•", "-", "·", "*")):
                    current_entry["description"].append(line.lstrip("•-·* "))
                elif len(line) > 30:
                    current_entry["description"].append(line)

        if current_entry:
            results.append(current_entry)

        return results

    def _parse_projects(self, proj_text: str) -> list:
        """Parse projects section."""
        results = []
        lines = [l.strip() for l in proj_text.split("\n") if l.strip()]

        current_project = None
        for line in lines:
            # Project names tend to be short title-like lines
            if len(line.split()) <= 6 and not line.startswith(("•", "-", "·")):
                if current_project:
                    results.append(current_project)
                current_project = {
                    "name": line,
                    "description": "",
                    "technologies": [],
                    "impact": None,
                    "url": None,
                }
                # Check for URL in same line
                url_match = re.search(r"https?://\S+", line)
                if url_match:
                    current_project["url"] = url_match.group()
            elif current_project:
                if not current_project["description"]:
                    current_project["description"] = line
                # Extract tech mentioned
                techs = [s.title() for s in self.TECH_SKILLS
                         if re.search(r"\b" + re.escape(s) + r"\b", line.lower())]
                current_project["technologies"].extend(techs)

        if current_project:
            results.append(current_project)

        return results

    def parse(self, file_bytes: bytes, filename: str) -> dict:
        """
        Main entry point: extract raw text, then parse all sections.
        Returns a dict matching the ParsedResume schema.
        """
        raw_text = self.extract_text(file_bytes, filename)
        sections = self._split_sections(raw_text)

        parsed = {
            "raw_text": raw_text,
            "name": self._extract_name(raw_text),
            "email": self._extract_email(raw_text),
            "phone": self._extract_phone(raw_text),
            "linkedin": self._extract_linkedin(raw_text),
            "github": self._extract_github(raw_text),
            "location": None,
            "summary": sections.get("summary", "").strip() or None,
            "skills": self._extract_skills(raw_text),
            "education": self._parse_education(sections.get("education", "")),
            "experience": self._parse_experience(sections.get("experience", "")),
            "projects": self._parse_projects(sections.get("projects", "")),
            "certifications": [],
            "languages": [],
        }

        # Extract location heuristically
        location_match = re.search(
            r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)*,\s*[A-Z]{2}(?:\s*\d{5})?)\b", raw_text
        )
        if location_match:
            parsed["location"] = location_match.group()

        # Extract certifications
        cert_text = sections.get("certifications", "")
        if cert_text:
            parsed["certifications"] = [
                l.strip().lstrip("•-·* ") for l in cert_text.split("\n")
                if l.strip() and len(l.strip()) > 5
            ]

        logger.info(f"Parsed resume: {parsed['name']}, {len(parsed['skills'])} skills")
        return parsed


# Singleton instance
resume_parser = ResumeParserService()
